from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException
from fastapi.responses import StreamingResponse
import io
import json
from .schemas import ScanRequest, ScanResponse
from .worker import process_scan, ACTIVE_SCANS_PROGRESS
import uuid
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy.orm import selectinload
from .database import get_db
from .models import ScanResult, Repository

router = APIRouter()

@router.get("/health")
async def health_check():
    return {"status": "ok", "message": "Backend is awake and healthy"}

@router.post("/scan", response_model=ScanResponse)
async def start_scan(request: ScanRequest, background_tasks: BackgroundTasks):
    task_id = str(uuid.uuid4())
    
    # Trigger scan as a background task directly
    background_tasks.add_task(process_scan, str(request.repo_url), task_id)
    
    return ScanResponse(
        status="success",
        message=f"Scan initiated for {request.repo_url}",
        task_id=task_id
    )

@router.get("/scans")
async def get_scans(db: AsyncSession = Depends(get_db)):
    result = await db.execute(
        select(ScanResult)
        .options(selectinload(ScanResult.repository), selectinload(ScanResult.findings))
        .order_by(ScanResult.created_at.desc())
    )
    scans = result.scalars().all()
    
    # Format for frontend
    response = []
    for scan in scans:
        response.append({
            "id": scan.id,
            "task_id": scan.task_id,
            "status": scan.status,
            "message": scan.message,
            "risk_score": scan.risk_score,
            "created_at": scan.created_at.isoformat() if scan.created_at else None,
            "repository": {
                "name": scan.repository.name,
                "url": scan.repository.url
            },
            "findings": [{"scanner_type": f.scanner_type, "raw_output": f.raw_output} for f in scan.findings]
        })
    return response

@router.get("/scan/{task_id}")
async def get_scan(task_id: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(
        select(ScanResult)
        .options(selectinload(ScanResult.repository), selectinload(ScanResult.findings))
        .where(ScanResult.task_id == task_id)
    )
    scan = result.scalars().first()
    if not scan:
        raise HTTPException(status_code=404, detail="Scan not found")
        
    return {
        "id": scan.id,
        "task_id": scan.task_id,
        "status": scan.status,
        "message": scan.message,
        "risk_score": scan.risk_score,
        "created_at": scan.created_at.isoformat() if scan.created_at else None,
        "repository": {
            "name": scan.repository.name,
            "url": scan.repository.url
        },
        "findings": [{"scanner_type": f.scanner_type, "raw_output": f.raw_output} for f in scan.findings]
    }

@router.get("/scan/{task_id}/progress")
async def get_progress(task_id: str):
    progress = ACTIVE_SCANS_PROGRESS.get(task_id, "Pending...")
    return {"task_id": task_id, "progress": progress}

@router.get("/scan/{task_id}/export/docx")
async def export_docx(task_id: str, db: AsyncSession = Depends(get_db)):
    from docx import Document
    from docx.shared import Pt, RGBColor
    
    result = await db.execute(
        select(ScanResult).options(selectinload(ScanResult.repository), selectinload(ScanResult.findings))
        .where(ScanResult.task_id == task_id)
    )
    scan = result.scalars().first()
    if not scan:
        raise HTTPException(status_code=404, detail="Scan not found")
        
    doc = Document()
    doc.add_heading(f'Security Scan Report: {scan.repository.name}', 0)
    doc.add_paragraph(f'Repository URL: {scan.repository.url}')
    doc.add_paragraph(f'Risk Score: {scan.risk_score}')
    doc.add_paragraph(f'Status: {scan.status.upper()}')
    
    doc.add_heading('Findings', level=1)
    
    for f in scan.findings:
        if f.scanner_type == 'semgrep':
            try:
                data = json.loads(f.raw_output)
                if data.get('results'):
                    doc.add_heading('Semgrep (SAST)', level=2)
                    for res in data['results']:
                        p = doc.add_paragraph()
                        run = p.add_run(f"Vulnerability: {res['check_id']}\n")
                        run.bold = True
                        p.add_run(f"File: {res['path']} (Line {res['start']['line']})\n")
                        p.add_run(f"Description: {res['extra']['message']}")
            except:
                pass
        elif f.scanner_type == 'trivy':
            try:
                data = json.loads(f.raw_output)
                if data.get('Results'):
                    doc.add_heading('Trivy (Dependencies & Legal)', level=2)
                    for res in data['Results']:
                        if res.get('Vulnerabilities'):
                            for v in res['Vulnerabilities']:
                                p = doc.add_paragraph()
                                run = p.add_run(f"CVE: {v['VulnerabilityID']} in {v['PkgName']}\n")
                                run.bold = True
                                p.add_run(f"Severity: {v['Severity']}\n")
                                p.add_run(f"Fix: Upgrade to {v.get('FixedVersion', 'N/A')}\n")
                                p.add_run(f"Description: {v.get('Description', '')}")
            except:
                pass
    
    f_stream = io.BytesIO()
    doc.save(f_stream)
    f_stream.seek(0)
    
    return StreamingResponse(
        f_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=report_{scan.repository.name}.docx"}
    )

